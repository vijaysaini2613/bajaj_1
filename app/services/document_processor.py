import asyncio
import httpx
import pypdf
import docx
import io
import os
from typing import List, Dict, Any
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse

from app.models.request_models import DocumentInput, DocumentType
from app.utils.text_processing import clean_text, extract_clauses
from app.utils.logger import setup_logger

logger = setup_logger(__name__)

class DocumentProcessor:
    """
    Service for processing various document types and extracting text content
    """
    
    def __init__(self):
        self.supported_extensions = {
            'pdf': self._process_pdf_content,
            'docx': self._process_docx_content,
            'doc': self._process_docx_content,
            'txt': self._process_text_content,
            'html': self._process_html_content,
            'htm': self._process_html_content
        }
        
        # Configure httpx client for URL downloads
        self.http_client = httpx.AsyncClient(
            timeout=30.0,
            follow_redirects=True,
            headers={
                'User-Agent': 'Document-QA-System/1.0'
            }
        )
    
    async def process_document(self, document: DocumentInput) -> List[str]:
        """
        Process a document and return extracted text chunks
        
        Args:
            document: Document input with type and content
            
        Returns:
            List of text chunks extracted from the document
        """
        try:
            logger.info(f"Processing document of type: {document.type}")
            
            if document.type == DocumentType.TEXT:
                return await self._process_text_document(document.content)
            elif document.type == DocumentType.URL:
                return await self._process_url_document(document.content)
            elif document.type == DocumentType.PDF:
                return await self._process_pdf_document(document.content)
            elif document.type == DocumentType.DOCX:
                return await self._process_docx_document(document.content)
            else:
                raise ValueError(f"Unsupported document type: {document.type}")
                
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            raise
    
    async def _process_text_document(self, content: str) -> List[str]:
        """Process plain text content"""
        cleaned_text = clean_text(content)
        if not cleaned_text:
            return []
        
        chunks = extract_clauses(cleaned_text)
        logger.info(f"Extracted {len(chunks)} chunks from text document")
        return chunks
    
    async def _process_url_document(self, url: str) -> List[str]:
        """Process document from URL"""
        try:
            logger.info(f"Downloading document from URL: {url}")
            
            async with self.http_client as client:
                response = await client.get(url)
                response.raise_for_status()
                
                content_type = response.headers.get('content-type', '').lower()
                content = response.content
                
                # Determine file type from content-type or URL extension
                if 'pdf' in content_type or url.lower().endswith('.pdf'):
                    return await self._process_pdf_content(content)
                elif 'msword' in content_type or 'wordprocessingml' in content_type or url.lower().endswith(('.docx', '.doc')):
                    return await self._process_docx_content(content)
                elif 'html' in content_type or url.lower().endswith(('.html', '.htm')):
                    return await self._process_html_content(content.decode('utf-8', errors='ignore'))
                else:
                    # Try as text
                    text_content = content.decode('utf-8', errors='ignore')
                    return await self._process_text_document(text_content)
                    
        except Exception as e:
            logger.error(f"Error processing URL document: {str(e)}")
            raise
    
    async def _process_pdf_document(self, content: str) -> List[str]:
        """Process PDF content (base64 encoded)"""
        try:
            import base64
            pdf_data = base64.b64decode(content)
            return await self._process_pdf_content(pdf_data)
        except Exception as e:
            logger.error(f"Error processing PDF document: {str(e)}")
            raise
    
    async def _process_docx_document(self, content: str) -> List[str]:
        """Process DOCX content (base64 encoded)"""
        try:
            import base64
            docx_data = base64.b64decode(content)
            return await self._process_docx_content(docx_data)
        except Exception as e:
            logger.error(f"Error processing DOCX document: {str(e)}")
            raise
    
    async def _process_pdf_content(self, content: bytes) -> List[str]:
        """Extract text from PDF content"""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = pypdf.PdfReader(pdf_file)
            
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num}: {str(e)}")
                    continue
            
            full_text = '\n\n'.join(text_parts)
            cleaned_text = clean_text(full_text)
            
            if not cleaned_text:
                logger.warning("No text could be extracted from PDF")
                return []
            
            chunks = extract_clauses(cleaned_text)
            logger.info(f"Extracted {len(chunks)} chunks from PDF document")
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing PDF content: {str(e)}")
            raise
    
    async def _process_docx_content(self, content: bytes) -> List[str]:
        """Extract text from DOCX content"""
        try:
            docx_file = io.BytesIO(content)
            document = docx.Document(docx_file)
            
            text_parts = []
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            full_text = '\n\n'.join(text_parts)
            cleaned_text = clean_text(full_text)
            
            if not cleaned_text:
                logger.warning("No text could be extracted from DOCX")
                return []
            
            chunks = extract_clauses(cleaned_text)
            logger.info(f"Extracted {len(chunks)} chunks from DOCX document")
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing DOCX content: {str(e)}")
            raise
    
    async def _process_html_content(self, content: str) -> List[str]:
        """Extract text from HTML content"""
        try:
            soup = BeautifulSoup(content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text
            text = soup.get_text()
            cleaned_text = clean_text(text)
            
            if not cleaned_text:
                logger.warning("No text could be extracted from HTML")
                return []
            
            chunks = extract_clauses(cleaned_text)
            logger.info(f"Extracted {len(chunks)} chunks from HTML document")
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing HTML content: {str(e)}")
            raise
    
    async def _process_text_content(self, content: str) -> List[str]:
        """Process plain text content"""
        cleaned_text = clean_text(content)
        if not cleaned_text:
            return []
        
        chunks = extract_clauses(cleaned_text)
        logger.info(f"Extracted {len(chunks)} chunks from text content")
        return chunks
    
    async def __aenter__(self):
        return self
    
    async def __aexit__(self, exc_type, exc_val, exc_tb):
        await self.http_client.aclose()
